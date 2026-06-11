#!/usr/bin/env python3
"""Generate real PDF and DOCX test fixtures for verifying the ingestion sidecar.

Dev-only (uses the `dev` dependency group: reportlab + python-docx). Run:

    uv run --group dev python make_fixtures.py

Writes into ./fixtures/ (gitignored):
  * simple.pdf   — clean single-column prose (MarkItDown fast path)
  * tables.pdf   — a document with a real table (Docling fallback territory)
  * simple.docx  — prose + a table in a Word document
"""

from __future__ import annotations

import os

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")

PROSE = (
    "The Sea Ranch (Condominium One) is a residential complex on the Sonoma "
    "County coast of California, built between 1963 and 1965. It was designed "
    "by the firm MLTW: Charles Moore, Donlyn Lyndon, William Turnbull, and "
    "Richard Whitaker. The design uses rough-sawn wood and shed roofs borrowed "
    "from local agricultural buildings, rejecting the International Style in "
    "favor of regional, site-specific architecture."
)


def make_simple_pdf() -> str:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    path = os.path.join(FIXTURES_DIR, "simple.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("The Sea Ranch", styles["Title"]),
        Spacer(1, 12),
        Paragraph(PROSE, styles["BodyText"]),
    ]
    doc.build(story)
    return path


def make_tables_pdf() -> str:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    path = os.path.join(FIXTURES_DIR, "tables.pdf")
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    data = [
        ["Building", "Architect", "Year"],
        ["Sea Ranch", "MLTW", "1965"],
        ["Vanna Venturi House", "Robert Venturi", "1964"],
        ["Guild House", "Venturi & Rauch", "1963"],
    ]
    table = Table(data)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
            ]
        )
    )
    story = [
        Paragraph("Key Works of Early Postmodern Architecture", styles["Title"]),
        Spacer(1, 12),
        table,
    ]
    doc.build(story)
    return path


def make_simple_docx() -> str:
    from docx import Document

    path = os.path.join(FIXTURES_DIR, "simple.docx")
    document = Document()
    document.add_heading("The Sea Ranch", level=1)
    document.add_paragraph(PROSE)
    document.add_heading("Key Works", level=2)
    table = document.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Building", "Architect", "Year"
    for building, architect, year in [
        ("Sea Ranch", "MLTW", "1965"),
        ("Vanna Venturi House", "Robert Venturi", "1964"),
    ]:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = building, architect, year
    document.save(path)
    return path


def main() -> None:
    os.makedirs(FIXTURES_DIR, exist_ok=True)
    for fn in (make_simple_pdf, make_tables_pdf, make_simple_docx):
        print(f"wrote {fn()}")


if __name__ == "__main__":
    main()
